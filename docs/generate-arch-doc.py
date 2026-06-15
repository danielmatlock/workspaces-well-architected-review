#!/usr/bin/env python3
"""Generate Architecture Design Document for WorkSpaces WAFR Tool"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Title Page ---
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Architecture Design Document')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(35, 47, 62)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('WorkSpaces Well-Architected Framework Review Tool')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(255, 153, 0)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('AWS Account: 590183747733 | Region: eu-west-2 (London)\n').font.size = Pt(11)
meta.add_run('Author: Daniel Matlock | Date: June 2026\n').font.size = Pt(11)
meta.add_run('Classification: Internal').font.size = Pt(11)

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Solution Overview',
    '3. Architecture Design',
    '4. Component Details',
    '5. Data Flow',
    '6. Security',
    '7. Backup & Recovery',
    '8. Cost Analysis',
    '9. Deployment Process',
    '10. Known Issues & Future Improvements',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# --- 1. Executive Summary ---
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document describes the architecture of the WorkSpaces Well-Architected Framework Review (WAFR) Tool — '
    'a cloud-hosted web application built for AWS Technical Account Managers (TAMs) to conduct structured '
    'Well-Architected reviews of Amazon WorkSpaces deployments.'
)
doc.add_paragraph(
    'The application is a single-page HTML application hosted on AWS Amplify, with authentication via Amazon Cognito, '
    'data persistence via AWS AppSync and Amazon DynamoDB, and AI-powered review guidance via Amazon Bedrock '
    '(Claude Haiku 4.5) accessed through API Gateway and Lambda.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Design Principles:').bold = True
bullets = [
    'Serverless-first — no infrastructure to manage, scales to zero when unused',
    'Cost-optimised — entire solution runs under $1/month for 1-2 users',
    'Single-page application — one HTML file, no build pipeline required',
    'Offline-capable — falls back to localStorage when cloud services unavailable',
    'AI-augmented — Bedrock provides on-demand review guidance without pre-generation',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# --- 2. Solution Overview ---
doc.add_heading('2. Solution Overview', level=1)

doc.add_heading('2.1 Purpose', level=2)
doc.add_paragraph(
    'The tool enables TAMs to conduct comprehensive Well-Architected Framework Reviews specifically tailored '
    'for Amazon WorkSpaces deployments. It covers 6 pillars (Operational Excellence, Security, Reliability, '
    'Performance Efficiency, Cost Optimisation, Sustainability) with 60+ questions covering WorkSpaces-specific '
    'best practices.'
)

doc.add_heading('2.2 Users', level=2)
doc.add_paragraph('Primary users: 1-2 AWS Technical Account Managers')
doc.add_paragraph('Access method: Web browser (any device with internet access)')
doc.add_paragraph('Authentication: Email + password via Amazon Cognito')

doc.add_heading('2.3 Technology Stack', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
hdr[2].text = 'Purpose'
rows = [
    ('Frontend', 'HTML + JavaScript (vanilla)', 'Single-page application'),
    ('Hosting', 'AWS Amplify + CloudFront', 'Static site delivery via CDN'),
    ('Authentication', 'Amazon Cognito + cognito-identity-js', 'User login (SRP auth)'),
    ('API (Data)', 'AWS AppSync (GraphQL)', 'CRUD operations for reviews/templates'),
    ('API (AI)', 'API Gateway + Lambda', 'AI guidance generation'),
    ('Database', 'Amazon DynamoDB', 'Reviews and templates storage'),
    ('AI', 'Amazon Bedrock (Claude Haiku 4.5)', 'On-demand review explanations'),
    ('Source Control', 'GitHub', 'Version control + deployment trigger'),
]
for layer, tech, purpose in rows:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = purpose

doc.add_page_break()

# --- 3. Architecture Design ---
doc.add_heading('3. Architecture Design', level=1)

doc.add_heading('3.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'The solution follows a serverless, event-driven architecture pattern. All compute is on-demand '
    '(Lambda), all storage is managed (DynamoDB), and all networking is handled by AWS managed services '
    '(CloudFront, API Gateway, AppSync).'
)

# ASCII diagram as a code block
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Architecture Diagram:').bold = True
doc.add_paragraph()

diagram_text = """
┌─────────────────────────────────────────────────────────────────────┐
│                        USER (Browser)                                │
└─────────────────────┬───────────────┬───────────────┬───────────────┘
                      │               │               │
                      ▼               ▼               ▼
              ┌──────────────┐ ┌─────────────┐ ┌──────────────┐
              │  CloudFront  │ │   AppSync   │ │ API Gateway  │
              │  (Amplify)   │ │  (GraphQL)  │ │ (HTTP API)   │
              └──────┬───────┘ └──────┬──────┘ └──────┬───────┘
                     │                │               │
                     ▼                ▼               ▼
              ┌──────────────┐ ┌─────────────┐ ┌──────────────┐
              │     S3       │ │  DynamoDB   │ │   Lambda     │
              │ (Static Site)│ │ (Data Store)│ │(wafr-explain)│
              └──────────────┘ └─────────────┘ └──────┬───────┘
                                                      │
                                                      ▼
                                               ┌──────────────┐
                                               │   Bedrock    │
                                               │(Claude Haiku)│
                                               └──────────────┘

              ┌──────────────┐
              │   Cognito    │ ◄── Authentication (JWT tokens)
              │ (User Pool)  │
              └──────────────┘
"""
para = doc.add_paragraph()
run = para.add_run(diagram_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_heading('3.2 Component Interactions', level=2)
doc.add_paragraph(
    'The browser loads the static HTML application from Amplify/CloudFront. On page load, '
    'the Cognito Identity JS SDK checks for an existing session. If authenticated, the app '
    'fetches reviews from AppSync (backed by DynamoDB). When a user views a question, '
    'the AI Guidance panel makes a POST request to API Gateway, which invokes Lambda, '
    'which calls Bedrock to generate contextual guidance.'
)

doc.add_heading('3.3 Design Decisions', level=2)
decisions = [
    ('Single HTML file', 'No build pipeline needed. Simplifies deployment and reduces moving parts. '
     'The entire app is one file that can be opened locally or served from any static host.'),
    ('Cognito Identity JS (not Amplify SDK)', 'Amplify v6 has no browser-compatible UMD bundle. '
     'The Cognito Identity JS SDK provides a working CDN-hosted bundle for SRP authentication.'),
    ('AppSync + DynamoDB (not direct DynamoDB)', 'AppSync provides GraphQL with built-in Cognito '
     'authorisation, eliminating the need for custom Lambda resolvers for basic CRUD.'),
    ('Separate Lambda for AI (not AppSync resolver)', 'Bedrock calls can take 5-10 seconds. '
     'A dedicated Lambda with 30s timeout avoids AppSync resolver timeout limits.'),
    ('EU inference profile for Bedrock', 'Data residency — keeps all AI processing within EU regions. '
     'Required for newer model versions that don\'t support direct on-demand invocation.'),
    ('Manual deployment (not CI/CD)', 'Amplify CI/CD has an unresolved IAM service role issue in this account. '
     'Manual zip upload via presigned URL works reliably.'),
]
for title_text, explanation in decisions:
    p = doc.add_paragraph()
    p.add_run(f'{title_text}: ').bold = True
    p.add_run(explanation)

doc.add_page_break()

# --- 4. Component Details ---
doc.add_heading('4. Component Details', level=1)

components = [
    ('4.1 Amplify Hosting', 'd1p2543h8l2mfc',
     'Serves the static HTML application via CloudFront CDN. Provides HTTPS, caching, and global edge delivery. '
     'Branch: main. Auto-build disabled (manual deployment).'),
    ('4.2 Amazon Cognito', 'eu-west-2_Wy0eJHyN3',
     'User Pool with email-based authentication. SRP (Secure Remote Password) protocol for login. '
     'Client ID: 1kj98mt2cjbci4sa9okfg4o5dk. No client secret (public web app). '
     'Password policy: 8+ chars, uppercase, lowercase, numbers required.'),
    ('4.3 AWS AppSync', '4up36qgqubd6tcuekx5cmexmii',
     'GraphQL API authenticated via Cognito User Pool tokens. Provides CRUD operations for reviews and templates. '
     'Uses VTL resolvers with $util.dynamodb.toMapValuesJson for automatic DynamoDB marshalling.'),
    ('4.4 Amazon DynamoDB', 'wafr-reviews, wafr-templates',
     'Two on-demand (PAY_PER_REQUEST) tables. Partition key: id (String). '
     'Point-in-Time Recovery enabled on both tables (35-day continuous backup). '
     'Reviews store: id, customerName, reviewerName, templateId, templateName, reviewDate, answers, notes, status.'),
    ('4.5 API Gateway', '6ylrfwa3d8',
     'HTTP API (not REST API — lower latency, lower cost). Single route: POST /explain. '
     'CORS configured for all origins. Auto-deploy stage. Timeout: 30 seconds.'),
    ('4.6 Lambda', 'wafr-explain',
     'Python 3.12 function. 128MB memory, 30s timeout. Receives question + best practice, '
     'calls Bedrock with a structured prompt, returns AI explanation. No VPC attachment needed.'),
    ('4.7 Amazon Bedrock', 'eu.anthropic.claude-haiku-4-5-20251001-v1:0',
     'EU inference profile for Claude Haiku 4.5. Fastest/cheapest Claude model. '
     'Max 512 tokens per response. Prompt optimised for concise, actionable TAM guidance. '
     'No fine-tuning or custom model — uses foundation model directly.'),
]

for heading, resource_id, description in components:
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    p.add_run(f'Resource: ').bold = True
    p.add_run(resource_id)
    doc.add_paragraph(description)

doc.add_page_break()

# --- 5. Data Flow ---
doc.add_heading('5. Data Flow', level=1)

doc.add_heading('5.1 Authentication Flow', level=2)
steps = [
    'User enters email + password in login form',
    'Browser calls Cognito via amazon-cognito-identity-js SDK (SRP protocol)',
    'Cognito validates credentials and returns ID Token (JWT)',
    'Token stored in browser localStorage by the SDK',
    'All subsequent API calls include the JWT in the Authorization header',
    'Token refreshes automatically via refresh token (30-day validity)',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('5.2 Data Persistence Flow', level=2)
steps = [
    'User creates/updates a review (answers questions, adds notes)',
    'App saves to localStorage immediately (instant, offline-capable)',
    'App calls AppSync createReview mutation with JWT auth',
    'AppSync resolver writes to DynamoDB wafr-reviews table',
    'On next login, app calls listReviews query to load from DynamoDB',
    'Cloud data merged with localStorage (cloud wins on conflict)',
    'Pillar structure attached from template (not stored in DynamoDB to save space)',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('5.3 AI Guidance Flow', level=2)
steps = [
    'User navigates to a question in the review',
    'Browser sends POST to API Gateway /explain endpoint',
    'Body contains: { question, bestPractice }',
    'API Gateway invokes Lambda function wafr-explain',
    'Lambda constructs prompt and calls Bedrock InvokeModel',
    'Bedrock generates explanation (typically 3-5 seconds)',
    'Lambda returns JSON response with explanation text',
    'Browser renders explanation in the AI Guidance panel',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# --- 6. Security ---
doc.add_heading('6. Security', level=1)

doc.add_heading('6.1 Authentication & Authorisation', level=2)
bullets = [
    'All API calls require a valid Cognito JWT token',
    'AppSync authorisation type: AMAZON_COGNITO_USER_POOLS',
    'API Gateway /explain endpoint is unauthenticated (public) — low risk as it only generates generic guidance',
    'No sensitive customer data passes through the AI endpoint',
    'Cognito password policy enforces complexity requirements',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.2 Data Protection', level=2)
bullets = [
    'All data encrypted in transit (HTTPS/TLS 1.2+)',
    'DynamoDB encryption at rest (AWS managed key)',
    'No customer PII stored — only reviewer name, customer org name, and assessment scores',
    'Bedrock data: not used for model training (AWS Bedrock data policy)',
    'EU inference profile ensures AI processing stays within EU regions',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.3 IAM Least Privilege', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'Permissions'
hdr[2].text = 'Scope'
iam_rows = [
    ('github-actions-amplify-deploy', 'sts:AssumeRoleWithWebIdentity', 'This repo + main branch only'),
    ('appsync-dynamodb-role', 'dynamodb:GetItem/PutItem/DeleteItem/UpdateItem/Query/Scan', 'wafr-reviews + wafr-templates tables only'),
    ('lambda-bedrock-role', 'bedrock:InvokeModel + Lambda basic execution', 'All Bedrock models (could be scoped further)'),
]
for role, perms, scope in iam_rows:
    row = table.add_row().cells
    row[0].text = role
    row[1].text = perms
    row[2].text = scope

doc.add_heading('6.4 Network Security', level=2)
bullets = [
    'No VPC — all services are AWS managed endpoints',
    'No inbound ports to manage',
    'CloudFront provides DDoS protection (AWS Shield Standard)',
    'API Gateway has default throttling (10,000 requests/second)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# --- 7. Backup & Recovery ---
doc.add_heading('7. Backup & Recovery', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Backup Method'
hdr[2].text = 'RPO'
hdr[3].text = 'RTO'
backup_rows = [
    ('DynamoDB (reviews)', 'Point-in-Time Recovery', '1 second', '<1 hour (table restore)'),
    ('DynamoDB (templates)', 'Point-in-Time Recovery', '1 second', '<1 hour (table restore)'),
    ('Application code', 'GitHub (version controlled)', '0 (every commit)', '<5 min (redeploy)'),
    ('Lambda code', 'GitHub + AWS versioning', '0', '<5 min (redeploy)'),
    ('Cognito users', 'No backup (1-2 users, recreate manually)', 'N/A', '<5 min'),
    ('AppSync schema', 'GitHub (schema.graphql)', '0', '<10 min (recreate)'),
]
for comp, method, rpo, rto in backup_rows:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = method
    row[2].text = rpo
    row[3].text = rto

doc.add_paragraph()
doc.add_paragraph(
    'Disaster Recovery Strategy: The entire infrastructure can be recreated from scratch using the '
    'commands documented in the CHANGELOG.md. DynamoDB data is protected by PITR. '
    'Total rebuild time from zero: approximately 30 minutes.'
)

doc.add_page_break()

# --- 8. Cost Analysis ---
doc.add_heading('8. Cost Analysis', level=1)

doc.add_paragraph(
    'The solution is designed to operate within or near the AWS Free Tier for the expected usage of 1-2 users '
    'conducting reviews a few times per month.'
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Expected Usage'
hdr[2].text = 'Free Tier'
hdr[3].text = 'Monthly Cost'
cost_rows = [
    ('Amplify Hosting', '<1GB served, <1000 builds', '15GB/month, 1000 min', '$0.00'),
    ('DynamoDB', '<100 reads/writes per day', '25 RCU/WCU free', '$0.00'),
    ('DynamoDB PITR', '2 tables, <1MB each', 'N/A', '~$0.20'),
    ('AppSync', '<1000 queries/month', '250K queries free', '$0.00'),
    ('Cognito', '1-2 monthly active users', '50,000 MAU free', '$0.00'),
    ('API Gateway', '<200 requests/month', '1M requests free', '$0.00'),
    ('Lambda', '<200 invocations, 128MB', '1M requests free', '$0.00'),
    ('Bedrock (Haiku 4.5)', '~100 calls, 512 tokens each', 'N/A', '~$0.05'),
    ('CloudWatch Logs', 'Lambda logs', '5GB free', '$0.00'),
    ('TOTAL', '', '', '<$1.00/month'),
]
for svc, usage, free, cost in cost_rows:
    row = table.add_row().cells
    row[0].text = svc
    row[1].text = usage
    row[2].text = free
    row[3].text = cost

doc.add_page_break()

# --- 9. Deployment Process ---
doc.add_heading('9. Deployment Process', level=1)

doc.add_heading('9.1 Current Process (Manual)', level=2)
steps = [
    'Make changes locally to src/index.html',
    'Build deployment zip: cd src && zip -r /tmp/wafr-deploy.zip .',
    'Commit and push to GitHub: git add -A && git commit -m "message" && git push origin main',
    'In CloudShell: aws amplify create-deployment --app-id d1p2543h8l2mfc --branch-name main --region eu-west-2',
    'From Mac: curl -T /tmp/wafr-deploy.zip \'<PRESIGNED_URL>\'',
    'In CloudShell: aws amplify start-deployment --app-id d1p2543h8l2mfc --branch-name main --job-id <ID> --region eu-west-2',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('9.2 Future State (Automated)', level=2)
doc.add_paragraph(
    'Once the Amplify CI/CD IAM service role issue is resolved, deployments will be automatic on push to main. '
    'The GitHub Actions workflow is already configured and paused — it just needs the Amplify build environment '
    'to successfully assume the service role.'
)

doc.add_page_break()

# --- 10. Known Issues ---
doc.add_heading('10. Known Issues & Future Improvements', level=1)

doc.add_heading('10.1 Known Issues', level=2)
issues = [
    'Amplify CI/CD not functional — IAM service role "Unable to assume" error persists regardless of trust policy configuration',
    'Claude 3 Haiku/Sonnet base model IDs deprecated — must use inference profile IDs',
    'amplify-service-role IAM role exists but is unused',
]
for issue in issues:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('10.2 Future Improvements', level=2)
improvements = [
    'Resolve Amplify CI/CD for automated deployments on git push',
    'Add S3 bucket for exported PDF/HTML reports',
    'Add template management via DynamoDB (currently only default template)',
    'Add multi-user support with shared templates (Cognito groups)',
    'Cache Bedrock responses in DynamoDB to avoid repeat calls for same question',
    'Add custom domain (e.g. wafr.internal.example.com)',
    'Consider WAF on API Gateway if exposed more broadly',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

# --- Save ---
output_path = '/Users/danmmat/Scripts/Github/workspaces-well-architected-review/docs/Architecture-Design-Document.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
